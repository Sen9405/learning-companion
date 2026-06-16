#!/usr/bin/env python3
"""Convert MY_ROADMAP.md to a structured .docx file."""

import os
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

SRC = os.path.expanduser("~/ai-engineer-agent-roadmap-2026/MY_ROADMAP.md")
DST = os.path.expanduser("~/ai-engineer-agent-roadmap-2026/Мой_путь_AI_Agent_Engineer_2026.docx")

# Colour palette
C_GREEN = RGBColor(0x22, 0xC5, 0x5E)
C_YELLOW = RGBColor(0xF5, 0xA6, 0x23)
C_ORANGE = RGBColor(0xE8, 0x6C, 0x00)
C_RED = RGBColor(0xDC, 0x35, 0x35)
C_DARK = RGBColor(0x1E, 0x1E, 0x2E)
C_MUTED = RGBColor(0x6B, 0x72, 0x80)
C_ACCENT = RGBColor(0x8B, 0x5C, 0xF6)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BG = RGBColor(0xF8, 0xF9, 0xFA)

PHASE_COLORS = {
    "Phase 0": C_GREEN,
    "Phase 1": C_YELLOW,
    "Phase 2": C_YELLOW,
    "Phase 3": C_ORANGE,
    "Phase 4": C_RED,
    "Phase 5": C_RED,
}

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = C_DARK
    h.font.name = 'Calibri'
    if level == 1:
        h.font.size = Pt(22)
    elif level == 2:
        h.font.size = Pt(16)
    else:
        h.font.size = Pt(13)


def add_phase_card(phase_title, body_text, color):
    """Add a phase as a styled section."""
    h = doc.add_heading(phase_title, level=2)
    for run in h.runs:
        run.font.color.rgb = color

    # Body
    for line in body_text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue

        # Bold markers
        is_bold = line.startswith("**") and "**" in line[2:]
        line_clean = line.replace("**", "")

        # Checkmark lines
        is_checkbox = line.startswith("- [ ]") or line.startswith("- [x]") or line.startswith("✅")
        is_bullet = line.startswith("- ") or line.startswith("· ")
        is_numbered = re.match(r"^\d+\.", line)
        is_link = line.startswith("[") and "](" in line

        p = doc.add_paragraph()
        p.space_after = Pt(4)
        p.space_before = Pt(2)

        if is_checkbox:
            p.style = doc.styles['Normal']
            run = p.add_run(f"☐ {line_clean.lstrip('- [ ]').lstrip('- [x]').lstrip('✅').strip()}")
            run.font.size = Pt(11)
            run.font.color.rgb = C_DARK
        elif is_bullet or is_numbered:
            prefix = "• " if is_bullet else ""
            text = line_clean.lstrip("- ").lstrip("· ")
            # Detect material links
            parts = re.split(r"(https?://\S+)", text)
            run = p.add_run(f"{prefix}")
            run.font.size = Pt(11)
            for part in parts:
                if part.startswith("http"):
                    r = p.add_run(part)
                    r.font.color.rgb = C_ACCENT
                    r.font.size = Pt(11)
                else:
                    r = p.add_run(part)
                    r.font.size = Pt(11)
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            if is_bold:
                for r in p.runs:
                    r.bold = True
        elif line.startswith(">"):
            text = line.lstrip(">").strip()
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = C_MUTED
            run.font.size = Pt(10)
        else:
            run = p.add_run(line_clean)
            run.font.size = Pt(11)

    # Separator
    doc.add_paragraph("─" * 50).runs[0].font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)


# =====================================
# TITLE PAGE
# =====================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("🧭 Мой путь\nAI Agent Engineer — 2026")
run.font.size = Pt(28)
run.font.color.rgb = C_DARK
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Персонализированный план обучения")
run.font.size = Pt(14)
run.font.color.rgb = C_MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("15 мая 2026")
run.font.size = Pt(11)
run.font.color.rgb = C_MUTED

doc.add_page_break()

# =====================================
# PROFILE
# =====================================
doc.add_heading("👤 Профиль", level=1)

profile_data = [
    ("Уровень", "Начинающий (vibe coding Telegram ботов)"),
    ("Время", "~15 часов/неделю"),
    ("Стек", "Python + DeepSeek / open-weights"),
    ("Цель", "Найти работу + запустить продукт"),
    ("Длительность", "~26 недель"),
]

table = doc.add_table(rows=len(profile_data), cols=2)
table.style = 'Light Shading Accent 1'
for i, (key, val) in enumerate(profile_data):
    cell_k = table.cell(i, 0)
    cell_v = table.cell(i, 1)
    cell_k.text = key
    cell_v.text = val
    for paragraph in cell_k.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(11)
    for paragraph in cell_v.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run("Расчёт: 17 недель (канон) × 1.5 (15 ч/нед) ≈ 26 недель")
run.font.size = Pt(9)
run.font.color.rgb = C_MUTED
run.italic = True

doc.add_paragraph()

# =====================================
# PHASES
# =====================================
doc.add_heading("📅 План по фазам", level=1)

phases_raw = """
Phase 0: Foundations — 3 недели
NORMAL — ты начинаешь, поэтому фундамент без спешки.

**Что добавил для тебя:**
- Перед стартом — HuggingFace LLM Course (токенизация, трансформеры — база)
- Open Models have crossed a threshold — open-source догоняет закрытые модели

**Основные материалы:**
- Building Effective Agents (Anthropic) — 5 паттернов, workflow vs agent
- Effective context engineering (Anthropic) — прочитать дважды
- Context Engineering for Agents (LangChain) — Write, Select, Compress, Isolate
- Anthropic Cookbook — код к каждому паттерну

**🎯 Проект:** Написать 2-страничный документ своими словами: workflow vs agent, augmented LLM, 4 контекстных примитива, orchestrator-worker, что такое harness.

✅ **Чекпоинт:** Можешь объяснить без терминов фреймворков, что такое агент.

---

Phase 1: Первый простой агент — 4 недели
NORMAL + время на привыкание к API.

**Что важно для тебя:**
- Используешь DeepSeek API (OpenAI-совместимый — tool use работает)
- Принципы те же: agent loop, tool_use, парсинг ответа модели

**Материалы:**
- Tutorial: Build a tool-using agent (Anthropic docs) — концепции универсальны
- Building agents with the Claude Agent SDK
- DeepSeek API docs — как устроен tool use

**🎯 Проект:** Написать tool-using агента дважды:
1. Сырой цикл ~100 строк: модель → парсинг tool_use → выполнение → результат
2. С использованием LiteLLM или готового SDK

✅ **Чекпоинт:** Агент, который может ходить в интернет/файлы и возвращать ответ.

---

Phase 2: Deep Agent (исследователь) — 5 недель
NORMAL. Самая объёмная фаза.

**Адаптация под open-weights:**
- LangGraph 1.0 работает с любыми моделями — используем с DeepSeek
- Вместо LangSmith → Phoenix (Arize) — open-source, бесплатно
- Deep Agents middleware от LangChain — тоже работает с любыми моделями

**Материалы:**
- LangGraph Quick Start
- LangChain Academy: Intro to LangGraph — бесплатно
- Deep Agents by LangChain — reference open-source harness
- Multi-agent research system (Anthropic) — orchestrator-worker

**🎯 Проект:** Исследователь-аналитик
- Агент получает вопрос → пишет план → запускает 3 саб-агента поиска → отчёт
- PostgresSaver (persistence) + human-in-the-loop на дорогих операциях
- Phoenix trace для одного полного прогона

✅ **Чекпоинт:** Работающий deep agent с durability и sub-agents.

---

Phase 3: Строим свой harness — 4 недели
SPEEDRUN — для трудоустройства достаточно понимания архитектуры. Для продукта — Deep Agents как база + модификации.

**Что меняем:**
- Не пишешь 1500 строк с нуля — форкаешь Deep Agents и добавляешь свой модуль
- Фокус на понимании, а не на копировании

**Материалы:**
- The Anatomy of an Agent Harness (LangChain)
- Improving Deep Agents with harness engineering
- deepagents source — читать исходники

**🎯 Проект:** Форкнуть Deep Agents + свой модуль (система промптов, хук pre_tool/post_tool) + post-mortem на 500+ слов.

✅ **Чекпоинт:** Post-mortem с сравнением форка и Claude Agent SDK.

---

Phase 4: Eval и CI — 5 недель
NORMAL + утяжелено (для поиска работы и продукта это ключ). Quality — #1 барьер в индустрии. За evals платят деньги.

**Материалы:**
- Demystifying evals for AI agents (Anthropic)
- Agent Evaluation Readiness Checklist (LangChain)
- Evaluating Deep Agents: Our Learnings
- Inspect AI (UK AISI) — benchmark-grade evals

**🎯 Проект:**
- Golden dataset 30-50 вопросов, 3 уровня сложности
- 4 типа eval'ов: single-turn, trajectory, LLM-as-judge, end-state
- CI-гейт в GitHub Actions
- Запуск benchmark через Inspect (GAIA Level 1)
- Бонус: опубликовать проект публично (GitHub + статья) — для портфолио

✅ **Чекпоинт:** make eval выдаёт CI pass/fail + Inspect лог.

---

Phase 5: Production Hardening — бессрочно
DEEP — для продукта самая важная фаза.

**Адаптация под open-weights:**
- Весь Phase 5 + отдельный фокус на cost-discipline
- Sandboxing: E2B (бесплатный tier) или Modal
- Credential broker — ключи API не попадают в контекст

**Deliverables:**
- Prompt caching + routing по сложности
- Model-routing с cost-per-task бюджетами
- Sandbox для кода
- Durable execution (Inngest / Temporal / LangGraph PostgresSaver)
- Trace sampling + drift alerts

✅ **Чекпоинт:** Агент, который переживает контакт с реальными пользователями.
"""

# Parse and render each phase
blocks = phases_raw.strip().split("\n---\n")
for block in blocks:
    lines = block.strip().split("\n")
    title = lines[0].strip()
    body = "\n".join(lines[1:]).strip()
    color = C_GREEN
    for ph, clr in PHASE_COLORS.items():
        if title.startswith(ph):
            color = clr
            break
    add_phase_card(title, body, color)

doc.add_page_break()

# =====================================
# RESOURCES
# =====================================
doc.add_heading("📚 Ресурсы", level=1)
resource_text = """Для начала (Phase 0):
- HuggingFace LLM Course — База: токенизация, трансформеры
- Building Effective Agents — Прочитать первым
- Effective Context Engineering — Самый важный текст
- Open Models Have Crossed a Threshold — Для open-weights пользователей
- Anthropic Cookbook — Код к паттернам

Бесплатные модели:
- DeepSeek (Deep Think / Chat) — лучший balance quality/cost
- GLM-5 / MiniMax M2.7 — догоняют закрытые frontier
- LiteLLM — прослойка для переключения провайдеров

На будущее:
- YouTube: Andrej Karpathy, AI Engineer канал
- Блоги: Anthropic Engineering Blog, LangChain Blog, Hamel Husain
- Рассылка: Latent Space
- Комьюнити: LangChain Discord, HuggingFace Discord"""

p = doc.add_paragraph()
run = p.add_run(resource_text.replace("- ", "• "))
run.font.size = Pt(11)

doc.add_page_break()

# =====================================
# CHECKLIST
# =====================================
doc.add_heading("📦 Чеклист прогресса", level=1)

checklist = [
    ("Phase 0 (нед 1–3)", "2-страничный документ с ментальными моделями"),
    ("Phase 1 (нед 4–7)", "Tool-using агент (сырой + на SDK)"),
    ("Phase 2 (нед 8–12)", "Deep Agent исследователь + Phoenix trace"),
    ("Phase 3 (нед 13–16)", "Форк Deep Agents + post-mortem"),
    ("Phase 4 (нед 17–21)", "Golden dataset + CI eval gate + Inspect"),
    ("Phase 5 (нед 22+)", "Production hardening"),
]

for phase, desc in checklist:
    p = doc.add_paragraph()
    run = p.add_run(f"☐ {phase}: {desc}")
    run.font.size = Pt(11)
    p.space_after = Pt(6)

# =====================================
# NEXT STEP
# =====================================
doc.add_paragraph()
doc.add_heading("🎯 Следующий шаг", level=2)
p = doc.add_paragraph()
run = p.add_run("Прочитать Building Effective Agents by Anthropic")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = C_ACCENT

p = doc.add_paragraph()
run = p.add_run("https://anthropic.com/research/building-effective-agents")
run.font.color.rgb = C_ACCENT
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run("15 минут чтения, с которых начинается весь путь. После прочтения напиши в Telegram «прочитал», и я скажу, что делать дальше.")
run.font.size = Pt(11)
run.font.color.rgb = C_DARK
run.italic = True

doc.save(DST)
print(f"OK → {DST}")
