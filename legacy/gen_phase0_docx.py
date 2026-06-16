#!/usr/bin/env python3
"""Generate the Phase 0 final document in DOCX format."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# Helper: add a heading with color
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)  # blue accent
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bold_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ── Title ──
title = doc.add_heading('Phase 0: Foundations', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = subtitle.add_run('Мой путь AI Agent Engineer — 2026')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════
# Section 1
# ═══════════════════════════════════════
add_heading('1. Workflow vs Agent', level=1)

add_bold_body('Workflow')
add_body(
    'Это когда мы жёстко прописываем цепочку шагов: сначала отправляем промпт → '
    'получаем ответ LLM → потом другой промпт с результатом → снова ответ. '
    'Каждый шаг предопределён кодом.'
)

add_bold_body('Agent')
add_body(
    'Это когда LLM сама решает, когда и какой инструмент вызвать, '
    'в каком порядке, и когда остановиться.'
)

add_bold_body('Правило')
add_body(
    'По умолчанию используем workflow. Переходим к агенту только когда '
    'workflow не хватает.'
)

# ═══════════════════════════════════════
# Section 2
# ═══════════════════════════════════════
add_heading('2. Augmented LLM', level=1)

add_body(
    'Это усиленная LLM — база для любого агента. Состоит из трёх компонентов:'
)

add_bullet(' — доступ к документам, базам данных', 'RAG (Retrieval)')
add_bullet(' — возможность вызывать функции (поиск, API, калькулятор)', 'Tools')
add_bullet(' — запоминать историю диалога', 'Memory')

add_body(
    'Без этого LLM — просто болванка. С этими тремя компонентами — '
    'основа для построения агента.'
)

# ═══════════════════════════════════════
# Section 3
# ═══════════════════════════════════════
add_heading('3. Пять паттернов построения агентов (Building Effective Agents)', level=1)

patterns = [
    ('Prompt Chaining', 'Шаг за шагом: результат A → вход B. Для задач, где каждый шаг важен сам по себе.'),
    ('Routing', 'LLM классифицирует запрос и направляет в нужный модуль.'),
    ('Parallelization', 'Запуск нескольких LLM одновременно для разных задач.'),
    ('Orchestrator-Workers', 'Центральный агент (оркестратор) управляет саб-агентами, каждый отвечает за свою функцию или роль.'),
    ('Evaluator-Optimizer', 'Одна LLM генерирует, другая оценивает и даёт фидбек на улучшение.'),
]

for i, (name, desc) in enumerate(patterns, 1):
    add_bullet(desc, f'{i}. {name}: ')

# ═══════════════════════════════════════
# Section 4
# ═══════════════════════════════════════
add_heading('4. Четыре контекстных примитива (Context Engineering)', level=1)

primitives = [
    ('Write', 'Писать качественные инструкции для LLM: персона, задача, формат вывода, стиль.'),
    ('Select', 'Выбирать только релевантную информацию в момент запроса, не тащить весь контекст (поиск по файлам, RAG).'),
    ('Compress', 'Сжимать контекст, когда он заполняется (суммаризация, фильтрация шума).'),
    ('Isolate', 'Изолировать контексты между разными частями системы / саб-агентами.'),
]

for name, desc in primitives:
    add_bullet(desc, f'{name} — ')

add_bold_body('Sandwich pattern')
add_body(
    'Самые важные инструкции — в начале (системный промпт) и в конце (формат вывода). '
    'Переменный контекст — в середине. Модель лучше запоминает начало и конец.'
)

# ═══════════════════════════════════════
# Section 5
# ═══════════════════════════════════════
add_heading('5. Состояние индустрии агентов (State of Agent Engineering, LangChain, 2026)', level=1)

add_body('Опрос 1300+ профессионалов. Ключевые цифры:')

stats = [
    '57.3% организаций уже имеют агентов в продакшене',
    'Качество — главный барьер (не стоимость, как в прошлом году)',
    '89% внедрили observability, 52% — evaluation',
    'Большинство использует несколько моделей, а не одну',
    'OpenAI — лидер, но Gemini, Claude и open-source тоже активно используются',
]
for s in stats:
    add_bullet(s)

add_bold_body('Топ-3 агентов в повседневной работе:')
agents_list = [
    'Coding agents (Claude Code, Cursor, Copilot, Windsurf)',
    'Research & Deep Research agents (ChatGPT, Claude, Gemini, Perplexity)',
    'Custom agents на LangChain / LangGraph под свои задачи',
]
for a in agents_list:
    add_bullet(a)

add_bold_body('Самый популярный use case:')
add_body(
    'Customer support (26.5%) — агенты ставят прямо перед клиентами. '
    'За ним research & data analysis (24.4%) и internal workflow automation (18%).'
)

add_bold_body('Тренд:')
add_body(
    'Индустрия движется от POC (прототипов) к продакшену. Вопрос уже не '
    '«стоит ли строить агентов», а «как внедрять их надёжно и в масштабе».'
)

# ═══════════════════════════════════════
# Section 6
# ═══════════════════════════════════════
add_heading('6. Что такое Harness (анатомия агентской обвязки)', level=1)

p = doc.add_paragraph()
run = p.add_run('Agent = Model + Harness')
run.bold = True
run.font.size = Pt(13)

add_body(
    'Если ты не модель — ты строишь harness. '
    'Harness — это весь код, конфигурация и логика исполнения, '
    'которые не являются самой моделью. Модель сама по себе — не агент. '
    'Агентом её делает harness.'
)

add_bold_body('Компоненты harness:')
components = [
    'System prompts — инструкции модели',
    'Tools / Skills / MCP — инструменты и их описания',
    'Инфраструктура — файловая система, песочница, браузер',
    'Логика оркестрации — спавн саб-агентов, передача задач, маршрутизация между моделями',
    'Hooks / Middleware — компакшн, продолжение работы, проверки',
]
for c in components:
    add_bullet(c)

add_bold_body('Проблемы, которые решает harness:')

add_body('1. Контекст (context rot) — модель тупеет, когда контекст переполняется.')
methods = [
    'Compaction: суммаризация контекста',
    'Tool call offloading: выгрузка больших выводов инструментов в файловую систему',
    'Skills: загрузка навыков по мере необходимости',
]
for m in methods:
    add_bullet(m)

add_body('2. Долгие задачи (long horizon) — агент должен работать автономно долго.')
long_methods = [
    'Файловая система + git для сохранения прогресса между сессиями',
    'Ralph Loop: хук перехватывает попытку модели завершить работу и запускает заново с чистым контекстом, подтягивая файлы из предыдущей итерации',
]
for m in long_methods:
    add_bullet(m)

add_body('3. Безопасность — песочницы для изолированного выполнения кода.')

add_bold_body('Будущее harness:')
future = [
    'Часть функционала harness будет впитываться в модели (планирование, долгая работа «из коробки»)',
    'Но harness engineering останется важным — правильная обвязка делает любую модель эффективнее',
    'Пример из статьи: смена только harness подняла кодинг-агента с Top 30 до Top 5 в Terminal Bench 2.0',
]
for f in future:
    add_bullet(f)

# ═══════════════════════════════════════
# Section 7
# ═══════════════════════════════════════
add_heading('7. Agent Engineering — определение', level=1)

p = doc.add_paragraph()
run = p.add_run(
    'Agent engineering — это итеративный процесс превращения LLM в надёжные системы. '
    'Поскольку агенты недетерминированы, инженерам нужно быстро итерировать, '
    'чтобы улучшать качество агентов.'
)
run.italic = True

# ═══════════════════════════════════════
# Save
# ═══════════════════════════════════════
output_dir = os.path.expanduser('~/ai-engineer-agent-roadmap-2026')
output_path = os.path.join(output_dir, 'Phase_0_Foundations.docx')
doc.save(output_path)
print(f'OK: {output_path} ({(os.path.getsize(output_path)/1024):.0f} KB)')
